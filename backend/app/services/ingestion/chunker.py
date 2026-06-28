"""
Unified ingestion pipeline supporting PDF, CSV, and DOCX.
Strictly rejects any other file formats. Import cost: zero (lazy imports).

Tables are extracted separately and chunked row-by-row to preserve
structure (column headers stay attached to each row's values) instead
of letting RecursiveCharacterTextSplitter break them mid-row.
"""
import csv
import os
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def process_file(file_path: str) -> list[Document]:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        from langchain_community.document_loaders import PyMuPDFLoader
        loader = PyMuPDFLoader(file_path)
        docs = loader.load()
        table_docs = _extract_pdf_tables(file_path)

    elif ext == ".csv":
        return _load_csv(file_path)

    elif ext == ".docx":
        docs = _load_docx(file_path)
        table_docs = _extract_docx_tables(file_path)

    else:
        raise ValueError(
            f"Unsupported file format. The system is configured to process only PDF, CSV, and DOCX files. "
            f"Rejected file: {os.path.basename(file_path)}"
        )

    chunker = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", "."],
        chunk_size=650,
        chunk_overlap=120,
    )
    text_chunks = chunker.split_documents(docs)

    return text_chunks + table_docs


def _load_csv(file_path: str) -> list[Document]:
    """
    Load CSV rows without LangChain loader dependencies. Each row becomes
    one chunk so table-like records stay intact and ingestion stays cheap.
    """
    table_docs: list[Document] = []

    try:
        rows = _read_csv_rows(file_path)
    except Exception:
        return table_docs

    if not rows:
        return table_docs

    headers = [cell.strip() for cell in rows[0]]
    data_rows = rows[1:] if len(rows) > 1 else rows
    has_headers = any(headers) and len(data_rows) > 0

    for row_index, row in enumerate(data_rows, start=1):
        values = [str(value).strip() for value in row]
        if not any(values):
            continue

        if has_headers:
            row_text = " | ".join(
                f"{_csv_header(headers, i)}: {value}"
                for i, value in enumerate(values)
                if value
            )
        else:
            row_text = " | ".join(value for value in values if value)

        if not row_text.strip():
            continue

        table_docs.append(
            Document(
                page_content=row_text,
                metadata={
                    "row": row_index,
                    "type": "table",
                    "source_type": "csv",
                },
            )
        )

    return table_docs


def _csv_header(headers: list[str], index: int) -> str:
    if index < len(headers) and headers[index]:
        return headers[index]
    return f"column_{index + 1}"


def _read_csv_rows(file_path: str) -> list[list[str]]:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            with open(file_path, newline="", encoding=encoding) as handle:
                sample = handle.read(2048)
                handle.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel
                return [row for row in csv.reader(handle, dialect=dialect)]
        except UnicodeDecodeError:
            continue
    with open(file_path, newline="", encoding="latin-1") as handle:
        return [row for row in csv.reader(handle)]


def _load_docx(file_path: str) -> list[Document]:
    """
    Load DOCX paragraphs through python-docx, avoiding docx2txt runtime
    drift in the Docker worker image.
    """
    from docx import Document as DocxDocument

    try:
        docx_file = DocxDocument(file_path)
    except Exception:
        return []

    paragraphs = [
        paragraph.text.strip()
        for paragraph in docx_file.paragraphs
        if paragraph.text.strip()
    ]

    if not paragraphs:
        return []

    return [
        Document(
            page_content="\n\n".join(paragraphs),
            metadata={"source_type": "docx"},
        )
    ]


def _extract_pdf_tables(file_path: str) -> list[Document]:
    """
    Extract tables from a PDF using camelot (lattice flavor — works best
    for tables with visible grid lines, like the W-4 withholding tables).
    Each table row becomes its own Document chunk so the embedding model
    sees complete, structured rows instead of broken fragments.
    """
    table_docs: list[Document] = []

    try:
        import camelot
        tables = camelot.read_pdf(file_path, pages="all", flavor="lattice")
    except Exception:
        return table_docs

    for table_index, table in enumerate(tables):
        df = table.df
        if df.empty:
            continue

        headers = df.iloc[0].tolist()
        page_number = table.page

        for row_index, row in df.iloc[1:].iterrows():
            row_text = " | ".join(
                f"{headers[i]}: {value}"
                for i, value in enumerate(row.tolist())
                if str(value).strip()
            )
            if not row_text.strip():
                continue

            table_docs.append(
                Document(
                    page_content=row_text,
                    metadata={
                        "page": page_number,
                        "type": "table",
                        "table_index": table_index,
                        "row_index": row_index,
                        "source_type": "pdf",
                    },
                )
            )

    return table_docs


def _extract_docx_tables(file_path: str) -> list[Document]:
    """
    Extract tables from a DOCX using python-docx. Each table row becomes
    its own Document chunk, mirroring the PDF table handling above.
    """
    from docx import Document as DocxDocument

    table_docs: list[Document] = []

    try:
        docx_file = DocxDocument(file_path)
    except Exception:
        return table_docs

    for table_index, table in enumerate(docx_file.tables):
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        rows = [row for row in rows if any(cell for cell in row)]
        if len(rows) < 2:
            continue

        headers = rows[0]

        for row_index, row in enumerate(rows[1:], start=1):
            row_text = " | ".join(
                f"{headers[i]}: {value}"
                for i, value in enumerate(row)
                if value.strip()
            )
            if not row_text.strip():
                continue

            table_docs.append(
                Document(
                    page_content=row_text,
                    metadata={
                        "section": f"table_{table_index}",
                        "type": "table",
                        "table_index": table_index,
                        "row_index": row_index,
                        "source_type": "docx",
                    },
                )
            )

    return table_docs
